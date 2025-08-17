"""Text extraction utilities for different document types."""

import logging
from io import BytesIO
from typing import List, Tuple, Optional
import re

try:
    from pypdf import PdfReader
except ImportError:
    from PyPDF2 import PdfReader

try:
    import docx
except ImportError:
    docx = None

logger = logging.getLogger(__name__)


def clean_text(text: str) -> str:
    """
    Clean and normalize extracted text.
    
    Args:
        text: Raw extracted text
    
    Returns:
        Cleaned text string
    """
    if not text:
        return ""
    
    # Remove null characters
    text = text.replace("\x00", " ")
    
    # Normalize whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove excessive newlines but preserve paragraph breaks
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
    
    # Clean up common OCR artifacts
    text = re.sub(r'[^\w\s\-.,;:!?()[\]{}"\'/\\@#$%^&*+=<>~`|]', ' ', text)
    
    return text.strip()


def extract_pdf_text(binary_content: bytes) -> List[Tuple[int, str]]:
    """
    Extract text from PDF file.
    
    Args:
        binary_content: PDF file content as bytes
    
    Returns:
        List of tuples containing (page_number, text_content)
    """
    pages = []
    
    try:
        reader = PdfReader(BytesIO(binary_content))
        
        for page_num, page in enumerate(reader.pages, start=1):
            try:
                text = page.extract_text() or ""
                cleaned_text = clean_text(text)
                
                if cleaned_text:  # Only add pages with actual content
                    pages.append((page_num, cleaned_text))
                    
            except Exception as e:
                logger.warning(f"Failed to extract text from page {page_num}: {e}")
                continue
                
    except Exception as e:
        logger.error(f"Failed to parse PDF: {e}")
        raise ValueError(f"Unable to parse PDF: {e}")
    
    return pages


def extract_docx_text(binary_content: bytes) -> str:
    """
    Extract text from DOCX file.
    
    Args:
        binary_content: DOCX file content as bytes
    
    Returns:
        Extracted text content
    """
    if docx is None:
        raise ImportError("python-docx package is required for DOCX extraction")
    
    try:
        document = docx.Document(BytesIO(binary_content))
        
        # Extract text from paragraphs
        paragraphs = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraphs.append(text)
        
        # Extract text from tables
        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        
        full_text = "\n".join(paragraphs)
        return clean_text(full_text)
        
    except Exception as e:
        logger.error(f"Failed to parse DOCX: {e}")
        raise ValueError(f"Unable to parse DOCX: {e}")


def extract_txt_text(binary_content: bytes) -> str:
    """
    Extract text from plain text file.
    
    Args:
        binary_content: Text file content as bytes
    
    Returns:
        Extracted text content
    """
    try:
        # Try different encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                text = binary_content.decode(encoding)
                return clean_text(text)
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail, use utf-8 with error handling
        text = binary_content.decode('utf-8', errors='replace')
        return clean_text(text)
        
    except Exception as e:
        logger.error(f"Failed to decode text file: {e}")
        raise ValueError(f"Unable to decode text file: {e}")


def detect_file_type(filename: str, content: bytes) -> str:
    """
    Detect file type based on filename and content.
    
    Args:
        filename: Original filename
        content: File content as bytes
    
    Returns:
        File type string ('pdf', 'docx', 'txt', 'unknown')
    """
    filename_lower = filename.lower()
    
    # Check by file extension first
    if filename_lower.endswith('.pdf'):
        return 'pdf'
    elif filename_lower.endswith('.docx'):
        return 'docx'
    elif filename_lower.endswith(('.txt', '.text')):
        return 'txt'
    
    # Check by file signature (magic bytes)
    if content.startswith(b'%PDF'):
        return 'pdf'
    elif content.startswith(b'PK\x03\x04') and b'word/' in content[:1024]:
        return 'docx'
    
    # Default to txt for other text-like content
    try:
        content[:1024].decode('utf-8')
        return 'txt'
    except UnicodeDecodeError:
        pass
    
    return 'unknown'


def extract_text_from_file(filename: str, binary_content: bytes) -> List[Tuple[Optional[int], str]]:
    """
    Extract text from a file based on its type.
    
    Args:
        filename: Original filename
        binary_content: File content as bytes
    
    Returns:
        List of tuples containing (page_number, text_content)
        For non-paginated formats, page_number will be None
    """
    file_type = detect_file_type(filename, binary_content)
    
    try:
        if file_type == 'pdf':
            return extract_pdf_text(binary_content)
        elif file_type == 'docx':
            text = extract_docx_text(binary_content)
            return [(None, text)] if text else []
        elif file_type == 'txt':
            text = extract_txt_text(binary_content)
            return [(None, text)] if text else []
        else:
            logger.warning(f"Unsupported file type '{file_type}' for file '{filename}'")
            return []
            
    except Exception as e:
        logger.error(f"Failed to extract text from {filename}: {e}")
        return []


def get_supported_extensions() -> List[str]:
    """
    Get list of supported file extensions.
    
    Returns:
        List of supported file extensions
    """
    return ['.pdf', '.docx', '.txt', '.text']
